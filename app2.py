from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from apscheduler.schedulers.background import BackgroundScheduler
from docx import Document
import io
from datetime import datetime
import csv
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

app = Flask(__name__)
app.config['SECRET_KEY'] = '1111'
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# Укажите абсолютный путь к файлу absents.csv
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CSV_FILE = os.path.join(BASE_DIR, 'absents.csv')


class User(UserMixin):
    def __init__(self, id, username, password):
        self.id = id
        self.username = username
        self.password = password


users = [User(id=1, username='admin', password='admin')]


@login_manager.user_loader
def load_user(user_id):
    return next((user for user in users if user.id == int(user_id)), None)


def initialize_csv():
    if not os.path.exists(CSV_FILE):
        with open(CSV_FILE, mode='w', newline='') as file:
            writer = csv.writer(file)
            writer.writerow(['name', 'details', 'date'])
        print("Initialized CSV file with headers: ['name', 'details', 'date']")


def save_to_csv(name, details, date):
    with open(CSV_FILE, mode='a', newline='') as file:
        writer = csv.writer(file)
        writer.writerow([name, details, date])
    print(f"Saved to CSV: name={name}, details={details}, date={date}")


def load_from_csv():
    if os.path.exists(CSV_FILE):
        with open(CSV_FILE, mode='r', encoding='utf-8', errors='replace') as file:
            reader = csv.DictReader(file)
            data = [{'name': row['name'], 'details': row['details'], 'date': row['date']} for row in reader]
            print(f"Loaded data from CSV: {data}")
            return data
    return []


def load_from_csv_for_today():
    today = datetime.now().date()
    absents_today = []
    if os.path.exists(CSV_FILE):
        with open(CSV_FILE, mode='r', encoding='utf-8', errors='replace') as file:
            reader = csv.DictReader(file)
            for row in reader:
                if 'date' in row:
                    date = datetime.strptime(row['date'], '%Y-%m-%d %H:%M:%S').date()
                    if date == today:
                        absents_today.append({'name': row['name'], 'details': row['details'], 'date': row['date']})
    print(f"Absents for today: {absents_today}")
    return absents_today


def send_email_with_attachment(to_email, subject, body, attachment_filename, attachment_data):
    from_email = 'your_gmail@example.com'
    from_password = 'your_app_password'

    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = to_email
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    part = MIMEBase('application', 'octet-stream')
    part.set_payload(attachment_data.read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', f'attachment; filename= {attachment_filename}')

    msg.attach(part)

    with smtplib.SMTP('smtp.gmail.com', 587) as server:
        server.starttls()
        server.login(from_email, from_password)
        server.sendmail(from_email, to_email, msg.as_string())
    print(f"Sent email to {to_email} with subject '{subject}'")


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/submit', methods=['POST'])
def submit():
    name = request.form['name']
    absent_persons = request.form.getlist('absent_person')
    reasons = request.form.getlist('reason')
    details = ', '.join([f"{person} ({reason})" for person, reason in zip(absent_persons, reasons)])
    date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    save_to_csv(name, details, date)
    return redirect(url_for('index'))


@app.route('/list')
@login_required
def list_absent():
    ordered_departments = [
        "Бухгалтерія", "Бюджетний відділ", "Розвиток", "Тендерний", "Обслуговування",
        "Відділ Продажу", "Фарм Отдел", "Маркетинг", "Зовнішня служба", "Відділ закупівель",
        "Фінансовий відділ", "Відділ IT", "Відділ персоналу", "Відділ контролю якості",
        "Відділ економ безпеки", "Керівники"
    ]
    absents_list = load_from_csv_for_today()
    ordered_absents = sorted(absents_list, key=lambda x: ordered_departments.index(x['name']) if x['name'] in ordered_departments else len(ordered_departments))
    return render_template('list.html', absents=ordered_absents)


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = next((user for user in users if user.username == username and user.password == password), None)
        if user:
            login_user(user)
            return redirect(url_for('list_absent'))
        else:
            flash('Login Unsuccessful. Please check username and password', 'danger')
    return render_template('login.html')


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))


@app.route('/download_doc')
@login_required
def download_doc():
    doc = Document()
    doc.add_heading('Список відсутніх', level=1)

    ordered_departments = [
        "Бухгалтерія", "Бюджетний відділ", "Розвиток", "Тендерний", "Обслуговування",
        "Відділ Продажу", "Фарм Отдел", "Маркетинг", "Зовнішня служба", "Відділ закупівель",
        "Фінансовий відділ", "Відділ IT", "Відділ персоналу", "Відділ контролю якості",
        "Відділ економ безпеки", "Керівники"
    ]

    absents_list = load_from_csv_for_today()
    print(f"Absents for today (before sorting): {absents_list}")
    ordered_absents = sorted(absents_list, key=lambda x: ordered_departments.index(x['name']) if x['name'] in ordered_departments else len(ordered_departments))
    print(f"Ordered absents for document: {ordered_absents}")

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Відділ'
    hdr_cells[2].text = 'Відсутній та причина'
    hdr_cells[3].text = 'Дата'

    counter = 1
    for absent in ordered_absents:
        print(f"Adding to document: {absent}")
        row_cells = table.add_row().cells
        row_cells[0].text = str(counter)
        name_run = row_cells[1].paragraphs[0].add_run(absent['name'])
        name_run.bold = True
        row_cells[2].text = absent['details']
        row_cells[3].text = absent['date']
        counter += 1

    f = io.BytesIO()
    doc.save(f)
    f.seek(0)

    return send_file(f, as_attachment=True, download_name='Список_відсутніх.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/send_email', methods=['POST'])
@login_required
def send_email():
    email = request.form['email']

    doc = Document()
    doc.add_heading('Список відсутніх', level=1)

    ordered_departments = [
        "Бухгалтерія", "Бюджетний відділ", "Розвиток", "Тендерний", "Обслуговування",
        "Відділ Продажу", "Фарм Отдел", "Маркетинг", "Зовнішня служба", "Відділ закупівель",
        "Фінансовий відділ", "Відділ IT", "Відділ персоналу", "Відділ контролю якості",
        "Відділ економ безпеки", "Керівники"
    ]

    absents_list = load_from_csv_for_today()
    ordered_absents = sorted(absents_list, key=lambda x: ordered_departments.index(x['name']) if x['name'] in ordered_departments else len(ordered_departments))
    print(f"Ordered absents for email document: {ordered_absents}")

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Відділ'
    hdr_cells[2].text = 'Відсутній та причина'
    hdr_cells[3].text = 'Дата'

    counter = 1
    for absent in ordered_absents:
        print(f"Adding to email document: {absent}")
        row_cells = table.add_row().cells
        row_cells[0].text = str(counter)
        name_run = row_cells[1].paragraphs[0].add_run(absent['name'])
        name_run.bold = True
        row_cells[2].text = absent['details']
        row_cells[3].text = absent['date']
        counter += 1

    f = io.BytesIO()
    doc.save(f)
    f.seek(0)

    send_email_with_attachment(email, 'Список відсутніх', 'Прилагаем список отсутствующих за сегодня.',
                               'Список_відсутніх.docx', f)

    flash('Email sent successfully!', 'success')
    return redirect(url_for('list_absent'))


if __name__ == '__main__':
    initialize_csv()
    scheduler = BackgroundScheduler()
    scheduler.add_job(lambda: None, 'cron', hour=0, minute=0)
    scheduler.start()
    app.run(debug=True, host='0.0.0.0')
